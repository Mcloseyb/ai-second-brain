"""
文档解析器 — 将外部文件转换为 Markdown
--------------------------------------
支持格式: .md / .txt / .docx / .pdf
每种格式有独立的解析器，统一返回 ParsedDocument。

设计原则:
  - 解析失败不抛异常，返回空内容 + 日志记录
  - 保留尽可能多的结构信息（标题、粗体、列表等）
  - 编码问题自动回退（UTF-8 → GBK）

使用方式:
    from app.core.document_parser import document_parser

    doc = await document_parser.parse_file("D:/notes/机器学习.docx")
    print(doc.title)    # "机器学习"
    print(doc.content)  # "# 第一章\n\n..."
"""

import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# 支持的扩展名
SUPPORTED_EXTENSIONS = {".md", ".txt", ".docx", ".pdf"}


# ============================================================
# 数据结构
# ============================================================

@dataclass
class ParsedDocument:
    """解析后的文档"""
    title: str
    content: str
    source_type: str          # "md" | "docx" | "pdf"
    word_count: int = 0
    warnings: list[str] = field(default_factory=list)

    def __post_init__(self):
        if not self.word_count and self.content:
            self.word_count = len(self.content.split())


class DocumentParseError(Exception):
    """文档解析异常"""
    def __init__(self, message: str, file_path: str | None = None):
        self.file_path = file_path
        super().__init__(message)


# ============================================================
# 文档解析器
# ============================================================

class DocumentParser:
    """将外部文件解析为 Markdown 文本"""

    # --------------------------------------------------------
    # 公共方法
    # --------------------------------------------------------

    @staticmethod
    async def parse_file(file_path: str | Path) -> ParsedDocument:
        """
        解析本地文件 → ParsedDocument

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument

        Raises:
            FileNotFoundError: 文件不存在
            DocumentParseError: 不支持的格式 / 解析失败
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise DocumentParseError(
                f"不支持的格式 '{ext}'，支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
                file_path=str(path),
            )

        title = path.stem  # 文件名（无扩展名）作为默认标题
        logger.info(f"解析文档: {path.name} ({ext})")

        try:
            if ext in (".md", ".txt"):
                content, warnings = DocumentParser._parse_text(path)
                source_type = "md"
            elif ext == ".docx":
                content, warnings = DocumentParser._parse_docx(path)
                source_type = "docx"
            elif ext == ".pdf":
                content, warnings = DocumentParser._parse_pdf(path)
                source_type = "pdf"
            else:
                raise DocumentParseError(f"未处理的格式: {ext}", file_path=str(path))
        except DocumentParseError:
            raise
        except Exception as e:
            logger.error(f"解析文档失败 — {path.name}: {e}")
            raise DocumentParseError(f"解析失败: {e}", file_path=str(path)) from e

        # 从内容中提取更好的标题（第一个 # 标题）
        extracted_title = DocumentParser._extract_title(content)
        if extracted_title:
            title = extracted_title

        return ParsedDocument(
            title=title,
            content=content,
            source_type=source_type,
            warnings=warnings,
        )

    @staticmethod
    async def parse_bytes(filename: str, data: bytes) -> ParsedDocument:
        """
        从内存中的字节数据解析文档（用于上传文件）

        Args:
            filename: 原始文件名（用于判断格式）
            data: 文件字节内容

        Returns:
            ParsedDocument
        """
        import tempfile

        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise DocumentParseError(
                f"不支持的格式 '{ext}'，支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
                file_path=filename,
            )

        # 写入临时文件后解析（docx/pdf 需要文件路径，不能从 bytes 直接读）
        suffix = ext if ext else ".tmp"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        try:
            return await DocumentParser.parse_file(tmp_path)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    # --------------------------------------------------------
    # 格式解析器（私有）
    # --------------------------------------------------------

    @staticmethod
    def _parse_text(path: Path) -> tuple[str, list[str]]:
        """
        解析纯文本文件 (.md / .txt)
        返回: (content, warnings)
        """
        warnings: list[str] = []

        # 尝试 UTF-8，失败则回退 GBK
        for encoding in ("utf-8", "gbk", "latin-1"):
            try:
                text = path.read_text(encoding=encoding)
                if encoding != "utf-8":
                    warnings.append(f"使用 {encoding.upper()} 编码读取（非 UTF-8）")
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        else:
            # 所有编码都失败，用 errors='replace' 强制读取
            text = path.read_text(encoding="utf-8", errors="replace")
            warnings.append("编码检测失败，已使用替换字符")

        if not text.strip():
            warnings.append("文件内容为空")

        # 标准化换行符
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        return text, warnings

    @staticmethod
    def _parse_docx(path: Path) -> tuple[str, list[str]]:
        """
        解析 Word 文档 (.docx) → Markdown

        处理规则:
          - Heading 1/2/3 style → # / ## / ###
          - Bold runs → **text**
          - Italic runs → *text*
          - 列表项（有 numPr 或有列表标记的段落） → - item
          - 其他 → 普通段落
        """
        warnings: list[str] = []
        from docx import Document

        doc = Document(str(path))
        md_lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                md_lines.append("")
                continue

            # 检测标题样式
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                level = 1
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    pass
                level = max(1, min(6, level))
                md_lines.append(f"{'#' * level} {text}")
                continue

            # 检测粗体/斜体
            line = ""
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                line += run_text

            # 如果 run 级别的处理结果为空，用原始文本
            display = line.strip() if line.strip() else text

            # 检测列表
            if para.style and para.style.name and "List" in para.style.name:
                md_lines.append(f"- {display}")
            else:
                md_lines.append(display)

        # 处理表格
        for table in doc.tables:
            md_lines.append("")
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    md_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

        content = "\n".join(md_lines)

        # 清理多余空行
        content = DocumentParser._clean_whitespace(content)

        if not content.strip():
            warnings.append("DOCX 解析后内容为空")

        return content, warnings

    @staticmethod
    def _parse_pdf(path: Path) -> tuple[str, list[str]]:
        """
        解析 PDF 文档 → Markdown

        使用 PyMuPDF (fitz) 逐页提取文本。
        尝试按字体大小推断标题（字号 > 14pt 的短行 → 标题）。
        """
        warnings: list[str] = []
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        md_lines: list[str] = []

        try:
            for page_num in range(len(doc)):
                page = doc[page_num]

                # 尝试获取结构化文本（含字体信息）
                blocks = page.get_text("dict").get("blocks", [])

                for block in blocks:
                    if block.get("type") != 0:  # 非文本块（图片等）
                        continue

                    block_lines = []
                    for line in block.get("lines", []):
                        line_text = ""
                        max_font_size = 0
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")
                            max_font_size = max(max_font_size, span.get("size", 0))

                        clean = line_text.strip()
                        if not clean:
                            continue

                        # 字体 > 14pt 且较短 → 可能是标题
                        if max_font_size >= 14 and len(clean) < 100:
                            clean = f"## {clean}"

                        block_lines.append(clean)

                    if block_lines:
                        md_lines.extend(block_lines)
                        md_lines.append("")  # 段落间空行

                # 页间分隔
                if page_num < len(doc) - 1:
                    md_lines.append("")
        finally:
            doc.close()

        content = "\n".join(md_lines)
        content = DocumentParser._clean_whitespace(content)

        if not content.strip():
            warnings.append("PDF 解析后内容为空（可能是扫描件/图片 PDF）")

        return content, warnings

    # --------------------------------------------------------
    # 工具方法
    # --------------------------------------------------------

    @staticmethod
    def _extract_title(content: str) -> Optional[str]:
        """从 Markdown 内容中提取第一个 # 标题"""
        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("# ") and len(stripped) > 2:
                return stripped[2:].strip()
        return None

    @staticmethod
    def _clean_whitespace(text: str) -> str:
        """清理多余空行：3+ 连续空行 → 2 空行"""
        import re
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


# 全局单例
document_parser = DocumentParser()
